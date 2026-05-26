from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/moon/TTCS")
TASK_DIR = ROOT / ".planning/quick/260522-oso-7-13-word"
ART = TASK_DIR / "artifacts"
OUT = ROOT / "第13组-团队任务协作管理系统-项目需求规格说明书.docx"
OLD_REPORT = Path("/Users/moon/Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files/wxid_exc4w07gh7q922_ebfe/msg/file/2026-05/实验7-第13组.docx")
EXP5_BUSINESS = ROOT / ".planning/quick/260521-bzs-5-word-docx/artifacts/business-use-case.png"
EXP5_SYSTEM = ROOT / ".planning/quick/260521-bzs-5-word-docx/artifacts/system-use-case.png"

FONT_REGULAR = "/System/Library/Fonts/STHeiti Light.ttc"
FONT_BOLD = "/System/Library/Fonts/STHeiti Medium.ttc"


def fnt(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def table_fixed(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_doc(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        s = styles[name]
        s.font.name = "Arial"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(5)


def p(doc, text="", style=None, bold=False, align=None):
    para = doc.add_paragraph(style=style)
    if text:
        run = para.add_run(text)
        run.bold = bold
    if align is not None:
        para.alignment = align
    return para


def add_caption(doc, text):
    para = p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in para.runs:
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(80, 80, 80)


def add_image(doc, image_path: Path, caption: str, width=6.65):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table_fixed(table, [2100, 7200])
    for r, (k, v) in enumerate(rows):
        table.cell(r, 0).text = k
        table.cell(r, 1).text = v
        set_cell_shading(table.cell(r, 0), "E8EEF5")
        for cell in table.rows[r].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
            if r == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    return table


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_fixed(table, widths)
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        set_cell_shading(table.cell(0, i), "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    table_fixed(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(9)
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    return table


def extract_original_media():
    out = ART / "original_media"
    out.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(OLD_REPORT) as z:
        for name in z.namelist():
            if name.startswith("word/media/") and Path(name).suffix.lower() in {".png", ".jpg", ".jpeg"}:
                (out / Path(name).name).write_bytes(z.read(name))
    return out / "image1.png", out / "image2.png"


def rounded(draw, xy, text, fill, outline=(51, 65, 85), size=30):
    draw.rounded_rectangle(xy, radius=24, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    box = draw.textbbox((0, 0), text, font=fnt(size, True))
    draw.text(((x1 + x2 - box[2]) / 2, (y1 + y2 - box[3]) / 2 - 4), text, font=fnt(size, True), fill=(15, 23, 42))


def arrow(draw, start, end, width=4):
    draw.line((start, end), fill=(51, 65, 85), width=width)
    import math
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    for a in (ang + 2.55, ang - 2.55):
        draw.line((end[0], end[1], end[0] + 24 * math.cos(a), end[1] + 24 * math.sin(a)), fill=(51, 65, 85), width=width)


def draw_activity_task(path: Path):
    img = Image.new("RGB", (1800, 1500), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "任务执行与进度更新活动图", font=fnt(44, True), fill=(15, 23, 42))
    lanes = [(70, 130, 560, 1420, "项目经理"), (560, 130, 1130, 1420, "项目成员"), (1130, 130, 1730, 1420, "系统")]
    for x1, y1, x2, y2, label in lanes:
        d.rectangle((x1, y1, x2, y2), outline=(148, 163, 184), width=3)
        d.rectangle((x1, y1, x2, y1 + 70), fill=(232, 238, 245), outline=(148, 163, 184), width=3)
        d.text((x1 + 20, y1 + 17), label, font=fnt(30, True), fill=(31, 77, 120))
    steps = [
        ((170, 250, 455, 340), "创建任务"),
        ((170, 470, 455, 560), "分配负责人"),
        ((680, 470, 990, 560), "接收任务"),
        ((680, 700, 990, 790), "更新状态"),
        ((680, 930, 990, 1020), "填写工作日志"),
        ((1240, 930, 1580, 1020), "同步进度数据"),
        ((1240, 1150, 1580, 1240), "生成进度报表"),
    ]
    for box, label in steps:
        rounded(d, box, label, (219, 234, 254))
    for s, e in [((312, 340), (312, 470)), ((455, 515), (680, 515)), ((835, 560), (835, 700)), ((835, 790), (835, 930)), ((990, 975), (1240, 975)), ((1410, 1020), (1410, 1150))]:
        arrow(d, s, e)
    d.ellipse((790, 1330, 880, 1420), fill=(15, 23, 42))
    arrow(d, (1410, 1240), (835, 1330))
    img.save(path)


def draw_activity_bug(path: Path):
    img = Image.new("RGB", (1800, 1500), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "Bug 提交与处理活动图", font=fnt(44, True), fill=(15, 23, 42))
    lanes = [(70, 130, 560, 1420, "测试人员"), (560, 130, 1130, 1420, "开发人员"), (1130, 130, 1730, 1420, "项目经理/系统")]
    for x1, y1, x2, y2, label in lanes:
        d.rectangle((x1, y1, x2, y2), outline=(148, 163, 184), width=3)
        d.rectangle((x1, y1, x2, y1 + 70), fill=(232, 238, 245), outline=(148, 163, 184), width=3)
        d.text((x1 + 20, y1 + 17), label, font=fnt(30, True), fill=(31, 77, 120))
    steps = [
        ((170, 250, 455, 340), "执行测试"),
        ((170, 470, 455, 560), "提交 Bug"),
        ((1240, 470, 1580, 560), "分配处理人"),
        ((680, 470, 990, 560), "定位问题"),
        ((680, 700, 990, 790), "修复并提交"),
        ((170, 700, 455, 790), "回归验证"),
        ((1240, 930, 1580, 1020), "关闭 Bug"),
    ]
    for box, label in steps:
        rounded(d, box, label, (220, 252, 231))
    for s, e in [((312, 340), (312, 470)), ((455, 515), (1240, 515)), ((1240, 515), (990, 515)), ((835, 560), (835, 700)), ((680, 745), (455, 745)), ((455, 745), (1240, 975))]:
        arrow(d, s, e)
    d.text((210, 845), "未通过则重新打开", font=fnt(24), fill=(100, 116, 139))
    arrow(d, (312, 790), (835, 790))
    img.save(path)


def write_report():
    ART.mkdir(parents=True, exist_ok=True)
    hierarchy_img, old_usecase_img = extract_original_media()
    task_activity = ART / "activity-task.png"
    bug_activity = ART / "activity-bug.png"
    draw_activity_task(task_activity)
    draw_activity_bug(bug_activity)

    doc = Document()
    style_doc(doc)

    title = p(doc, "第13组-团队任务协作管理系统-项目需求规格说明书", align=WD_ALIGN_PARAGRAPH.CENTER)
    title.runs[0].font.size = Pt(22)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor(15, 23, 42)
    p(doc, "实验7：编写需求规格说明书", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc)
    add_kv_table(doc, [
        ("项目名称", "团队任务协作管理系统（TTCS）"),
        ("小组", "第 13 组"),
        ("文档版本", "V1.0"),
        ("文档性质", "课程实验提交版"),
        ("编写日期", "2026 年 5 月 22 日"),
        ("适用对象", "课程指导教师、项目组成员、开发人员、测试人员"),
    ])
    doc.add_page_break()

    p(doc, "目录", "Heading 1")
    toc = [
        "1 引言",
        "2 需求概述",
        "3 功能需求",
        "4 非功能需求",
        "5 故障处理",
        "6 其他需求",
        "附录 A 需求来源说明",
    ]
    for item in toc:
        p(doc, item)
    doc.add_page_break()

    p(doc, "1 引言", "Heading 1")
    p(doc, "1.1 编写目的", "Heading 2")
    p(doc, "本文档用于明确团队任务协作管理系统（TTCS）的软件需求，统一系统功能、行为流程、运行环境、非功能质量和故障处理要求，为后续概要设计、详细设计、编码实现、测试验收和课程答辩提供依据。")
    p(doc, "1.2 项目背景", "Heading 2")
    p(doc, "小型研发团队在任务分配、执行记录、阻塞反馈、项目进度统计和代码评审方面常依赖即时通讯、表格或口头同步，容易出现信息分散、责任不清、进度不可追踪等问题。TTCS 面向 5-15 人小型团队，提供任务协作、看板流转、工作日志、阻塞上报、报表统计和代码审核等功能，提升协作透明度。")
    p(doc, "1.3 定义", "Heading 2")
    add_matrix(doc, ["术语", "说明"], [
        ("TTCS", "Team Task Collaboration System，团队任务协作管理系统。"),
        ("看板", "以任务状态列展示任务流转过程的可视化管理方式。"),
        ("工作日志", "项目成员围绕任务填写的工作内容、工时和进展记录。"),
        ("阻塞任务", "因依赖、资源、需求或技术问题暂时无法推进的任务。"),
        ("RBAC", "Role-Based Access Control，基于角色的访问控制。"),
        ("PR/MR", "代码合并请求，用于代码审核和协作开发。"),
    ], [1800, 7560])
    p(doc, "1.4 参考资料", "Heading 2")
    add_matrix(doc, ["编号", "资料名称", "用途"], [
        ("R1", "实验5 系统需求建模", "功能建模、参与者识别和用例模型来源。"),
        ("R2", "实验6 系统动态建模", "活动流程和行为模型来源。"),
        ("R3", "软件工程课程实验要求", "需求规格说明书结构和提交要求。"),
        ("R4", "团队任务协作管理系统项目资料", "系统目标、业务范围和功能边界。"),
    ], [900, 3960, 4500])

    p(doc, "2 需求概述", "Heading 1")
    p(doc, "2.1 系统目标", "Heading 2")
    add_matrix(doc, ["目标编号", "目标描述"], [
        ("G-01", "为小型研发团队提供任务、项目、成员、协作和统计的一体化 Web 平台。"),
        ("G-02", "支持任务从创建、分配、执行、流转、日志记录到完成统计的闭环管理。"),
        ("G-03", "通过阻塞上报、评论提醒和报表统计提升项目进度透明度。"),
        ("G-04", "为课程演示和后续团队自用扩展保留清晰的功能边界和接口设计基础。"),
    ], [1500, 7860])
    p(doc, "2.2 运行环境", "Heading 2")
    add_matrix(doc, ["类别", "要求"], [
        ("服务端", "Linux Ubuntu 22.04 或同等环境，Python/FastAPI，MySQL 8.0+，Redis 7+。"),
        ("客户端", "Chrome 90+、Edge 90+、Firefox 88+、Safari 14+。"),
        ("系统架构", "B/S 架构，前后端分离，支持响应式 Web 页面。"),
        ("网络环境", "校园网或互联网环境，正式部署时支持 HTTPS。"),
    ], [1800, 7560])
    p(doc, "2.3 条件与限制", "Heading 2")
    add_matrix(doc, ["限制编号", "限制说明"], [
        ("C-01", "一期仅实现 Web 端，不开发移动端 App。"),
        ("C-02", "一期以中文界面为主，多语言支持放入后续扩展。"),
        ("C-03", "一期账号体系以邮箱密码登录为主，第三方登录不作为核心范围。"),
        ("C-04", "系统用户规模按 100 人以内设计，面向课程和小团队场景。"),
        ("C-05", "核心任务协作和项目管理优先，复杂自动化分析能力放入二期。"),
    ], [1500, 7860])

    p(doc, "3 功能需求", "Heading 1")
    p(doc, "本章整合实验 5 的用例模型和实验 6 的活动模型，对 TTCS 的功能范围、参与者、用例和核心业务流程进行说明。",)
    p(doc, "3.1 功能层次模型", "Heading 2")
    p(doc, "系统功能可划分为任务管理、沟通协作和数据统计三类核心子系统。")
    add_image(doc, hierarchy_img, "图 3-1 功能层次结构图", 6.7)

    p(doc, "3.2 系统用例模型", "Heading 2")
    p(doc, "系统主要参与者包括访客、普通用户、团队管理员、项目经理、项目成员和系统管理员。核心用例覆盖账号管理、团队管理、项目管理、任务管理、日志与阻塞、协作沟通、报表统计和系统管理。")
    if EXP5_BUSINESS.exists():
        add_image(doc, EXP5_BUSINESS, "图 3-2 业务用例图", 6.8)
    if EXP5_SYSTEM.exists():
        add_image(doc, EXP5_SYSTEM, "图 3-3 系统用例图", 6.8)
    else:
        add_image(doc, old_usecase_img, "图 3-3 系统用例图", 6.8)

    p(doc, "3.3 活动模型", "Heading 2")
    p(doc, "核心活动流程用于描述任务协作中的动态行为，重点覆盖任务执行与进度更新、Bug 提交与处理两个流程。")
    add_image(doc, task_activity, "图 3-4 任务执行与进度更新活动图", 6.7)
    add_image(doc, bug_activity, "图 3-5 Bug 提交与处理活动图", 6.7)

    p(doc, "3.4 功能需求清单", "Heading 2")
    add_matrix(doc, ["编号", "模块", "需求描述", "优先级"], [
        ("FR-01", "账号管理", "访客可注册账号，用户可登录、找回密码和维护个人资料。", "高"),
        ("FR-02", "团队管理", "用户可创建团队，团队管理员可邀请成员、移除成员和变更角色。", "高"),
        ("FR-03", "项目管理", "项目经理可创建项目、配置项目看板、分配项目成员并查看项目进度。", "高"),
        ("FR-04", "任务管理", "项目成员可新建、编辑、删除任务，拖拽任务状态并添加子任务。", "高"),
        ("FR-05", "工作日志", "项目成员可记录工作内容、工时、日期和任务进展。", "高"),
        ("FR-06", "阻塞上报", "成员可在任务受阻时提交阻塞原因，系统通知负责人并进入阻塞清单。", "高"),
        ("FR-07", "协作沟通", "成员可在任务下评论、@提醒、上传文档并关联代码审核。", "中"),
        ("FR-08", "报表统计", "系统可生成项目进度、工时统计、阻塞汇总和任务完成情况报表。", "中"),
        ("FR-09", "代码审核", "系统支持绑定代码仓库、查看提交记录、发起代码审核和添加文件评论。", "中"),
        ("FR-10", "系统管理", "系统管理员可进行权限管理、操作日志查看、数据备份、系统参数和安全策略配置。", "高"),
    ], [900, 1500, 6000, 960])

    p(doc, "3.5 用例文档", "Heading 2")
    add_matrix(doc, ["用例编号", "用例名称", "参与者", "前置条件", "基本流程", "后置条件"], [
        ("UC-01", "用户管理", "访客、普通用户、系统管理员", "用户访问系统或已登录", "注册账号；登录系统；维护个人资料；管理员维护账号状态。", "用户账号和权限信息被保存。"),
        ("UC-02", "团队管理", "普通用户、团队管理员", "用户已登录", "创建团队；邀请成员；接受邀请；变更角色或移除成员。", "团队成员结构更新。"),
        ("UC-03", "项目管理", "团队管理员、项目经理", "团队已存在", "创建项目；填写项目信息；配置看板；添加项目成员。", "项目创建完成并可进行任务管理。"),
        ("UC-04", "任务分配", "项目经理、项目成员", "项目已创建且成员已加入", "创建任务；填写标题、描述、优先级和截止日期；指定负责人；系统发送通知。", "任务进入待办列表。"),
        ("UC-05", "进度跟踪", "项目经理、项目成员", "任务已分配", "成员更新任务状态和工作日志；系统同步进度；项目经理查看整体进度。", "任务状态和项目进度更新。"),
        ("UC-06", "Bug 管理", "测试人员、开发人员、项目经理", "测试发现问题或任务存在缺陷", "测试人员提交 Bug；项目经理分配；开发人员修复；测试人员回归验证。", "Bug 被关闭或重新打开。"),
        ("UC-07", "报表生成", "项目经理、系统管理员", "系统已有任务、工时、阻塞或 Bug 数据", "选择报表类型和时间范围；系统汇总数据；展示或导出报表。", "报表生成完成。"),
        ("UC-08", "系统维护", "系统管理员", "管理员已登录", "配置权限；查看操作日志；执行数据备份；维护安全策略。", "系统配置和安全记录更新。"),
    ], [900, 1350, 1650, 1800, 2700, 960])

    p(doc, "4 非功能需求", "Heading 1")
    add_matrix(doc, ["类别", "需求编号", "需求内容", "验收标准"], [
        ("性能", "NFR-01", "页面加载和接口响应应满足课程演示和小团队使用要求。", "常用页面 2 秒内完成加载，95% 接口响应小于 500ms。"),
        ("并发", "NFR-02", "系统应支持小团队同时在线操作。", "支持 100 个以内并发在线用户。"),
        ("安全", "NFR-03", "系统应保护账号、接口和上传文件安全。", "密码加密存储，JWT 鉴权，防 SQL 注入、XSS、CSRF。"),
        ("可靠性", "NFR-04", "系统应具备基础容错和恢复能力。", "服务异常有友好提示，关键数据每日备份。"),
        ("易用性", "NFR-05", "普通成员应能快速掌握核心流程。", "新用户 30 分钟内能完成任务查看、日志填写和状态流转。"),
        ("兼容性", "NFR-06", "系统应兼容主流桌面浏览器。", "Chrome、Edge、Firefox、Safari 主流版本可正常使用。"),
        ("可维护性", "NFR-07", "系统代码和文档应便于后续扩展。", "模块边界清晰，日志完整，接口和需求编号可追踪。"),
    ], [1200, 1100, 4200, 2860])

    p(doc, "5 故障处理", "Heading 1")
    add_matrix(doc, ["故障场景", "处理方式"], [
        ("网络异常", "保存用户已输入内容，提示稍后重试，必要时自动重新提交。"),
        ("登录失败", "提示账号或密码错误，连续多次失败后临时锁定账号并记录日志。"),
        ("服务异常", "前端展示友好错误信息，后端记录异常日志，管理员可查看。"),
        ("数据异常", "依赖每日备份和恢复机制，在限定时间内恢复关键数据。"),
        ("文件上传失败", "提示文件类型、大小或网络错误，允许用户重新上传。"),
        ("阻塞上报失败", "保留填写内容，提示用户重试，避免阻塞原因丢失。"),
    ], [2200, 7160])

    p(doc, "6 其他需求", "Heading 1")
    add_matrix(doc, ["类别", "说明"], [
        ("可扩展性", "后续可扩展移动端、多语言、AI 辅助评审和更复杂的代码分析能力。"),
        ("实时性", "通知、在线状态和看板状态流转应尽量实时反馈。"),
        ("可视化", "报表应支持进度条、燃尽图、工时统计和阻塞趋势展示。"),
        ("课程演示", "一期功能应优先保证核心协作闭环可演示、可讲解、可验收。"),
    ], [1800, 7560])

    p(doc, "附录 A 需求来源说明", "Heading 1")
    add_matrix(doc, ["来源", "整合方式"], [
        ("实验5：系统需求建模", "提取参与者、系统用例图和用例描述，形成第 3 章功能需求和用例文档。"),
        ("实验6：系统动态建模", "提取任务执行、进度更新、Bug 处理等活动流程，形成第 3.3 节活动模型。"),
        ("小组会议讨论", "确定性能、安全、可靠性、易用性、兼容性和可维护性等非功能需求。"),
    ], [2200, 7160])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "第13组 - TTCS 需求规格说明书"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)

    doc.save(OUT)
    return OUT


def verify_docx(path: Path):
    doc = Document(path)
    with zipfile.ZipFile(path) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
        xml = z.read("word/document.xml")
    print(path)
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} images={len(doc.inline_shapes)} media={len(media)} bytes={path.stat().st_size}")
    print(f"document_xml_bytes={len(xml)}")


if __name__ == "__main__":
    out = write_report()
    verify_docx(out)
