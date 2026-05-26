from __future__ import annotations

import math
import shutil
import textwrap
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/moon/TTCS")
SRC = Path("/Users/moon/Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files/wxid_exc4w07gh7q922_ebfe/msg/file/2026-05/第 13 组.docx")
TASK_DIR = ROOT / ".planning/quick/260521-bzs-5-word-docx"
ARTIFACTS = TASK_DIR / "artifacts"
OUT = ROOT / "第13组-团队任务协作管理系统-项目需求模型-补全用例图.docx"

FONT_REGULAR = "/System/Library/Fonts/STHeiti Light.ttc"
FONT_BOLD = "/System/Library/Fonts/STHeiti Medium.ttc"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont) -> tuple[int, int]:
    box = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=8)
    return box[2] - box[0], box[3] - box[1]


def wrap_label(label: str, width: int = 7) -> str:
    if len(label) <= width:
        return label
    return "\n".join(textwrap.wrap(label, width=width, break_long_words=True))


def draw_actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str, fnt: ImageFont.FreeTypeFont) -> tuple[int, int]:
    stroke = (49, 65, 88)
    draw.ellipse((x - 18, y - 55, x + 18, y - 19), outline=stroke, width=4)
    draw.line((x, y - 18, x, y + 35), fill=stroke, width=4)
    draw.line((x - 35, y, x + 35, y), fill=stroke, width=4)
    draw.line((x, y + 35, x - 32, y + 78), fill=stroke, width=4)
    draw.line((x, y + 35, x + 32, y + 78), fill=stroke, width=4)
    w, h = text_size(draw, label, fnt)
    draw.text((x - w / 2, y + 92), label, fill=stroke, font=fnt)
    return x, y


def draw_oval(draw: ImageDraw.ImageDraw, center: tuple[int, int], label: str, fill: tuple[int, int, int],
              width: int = 250, height: int = 88) -> tuple[int, int, int, int]:
    x, y = center
    box = (x - width // 2, y - height // 2, x + width // 2, y + height // 2)
    draw.ellipse(box, fill=fill, outline=(61, 89, 132), width=3)
    fnt = font(26)
    text = wrap_label(label)
    tw, th = text_size(draw, text, fnt)
    draw.multiline_text((x - tw / 2, y - th / 2 - 2), text, fill=(18, 33, 56), font=fnt, align="center", spacing=8)
    return box


def nearest_point_on_oval(box: tuple[int, int, int, int], from_xy: tuple[int, int]) -> tuple[int, int]:
    left, top, right, bottom = box
    cx = (left + right) / 2
    cy = (top + bottom) / 2
    rx = (right - left) / 2
    ry = (bottom - top) / 2
    dx = from_xy[0] - cx
    dy = from_xy[1] - cy
    if dx == 0 and dy == 0:
        return int(cx), int(cy)
    scale = math.sqrt((dx * dx) / (rx * rx) + (dy * dy) / (ry * ry))
    return int(cx + dx / scale), int(cy + dy / scale)


def line_to_oval(draw: ImageDraw.ImageDraw, start: tuple[int, int], oval_box: tuple[int, int, int, int],
                 color=(90, 111, 140), width: int = 3):
    end = nearest_point_on_oval(oval_box, start)
    draw.line((start, end), fill=color, width=width)


def dashed_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str = "", label_shift: tuple[int, int] = (0, 0)):
    color = (80, 96, 120)
    dash = 16
    gap = 10
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    dist = math.hypot(dx, dy)
    if dist == 0:
        return
    ux, uy = dx / dist, dy / dist
    t = 0
    while t < dist:
        t2 = min(t + dash, dist)
        draw.line((start[0] + ux * t, start[1] + uy * t, start[0] + ux * t2, start[1] + uy * t2), fill=color, width=3)
        t += dash + gap
    angle = math.atan2(dy, dx)
    for a in (angle + math.pi * 0.82, angle - math.pi * 0.82):
        draw.line((end[0], end[1], end[0] + 22 * math.cos(a), end[1] + 22 * math.sin(a)), fill=color, width=3)
    if not label:
        return
    lf = font(22)
    mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
    mx += label_shift[0]
    my += label_shift[1]
    tw, th = text_size(draw, label, lf)
    draw.rounded_rectangle((mx - tw / 2 - 10, my - th / 2 - 8, mx + tw / 2 + 10, my + th / 2 + 8), radius=8, fill=(255, 255, 255), outline=(203, 213, 225))
    draw.text((mx - tw / 2, my - th / 2), label, fill=color, font=lf)


def title(draw: ImageDraw.ImageDraw, text: str):
    draw.text((90, 48), text, fill=(15, 23, 42), font=font(44, bold=True))


def draw_business_diagram(path: Path):
    img = Image.new("RGB", (2600, 1550), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title(draw, "业务用例图")
    draw.rounded_rectangle((420, 150, 2180, 1430), radius=30, fill=(255, 255, 255), outline=(99, 102, 241), width=5)
    draw.text((470, 185), "团队任务协作管理系统", fill=(49, 46, 129), font=font(34, bold=True))

    actor_font = font(27, bold=True)
    actors = {
        "访客": (200, 360),
        "普通用户": (200, 720),
        "项目成员": (200, 1080),
        "团队管理员": (2380, 420),
        "项目经理": (2380, 790),
        "系统管理员": (2380, 1160),
    }
    for name, pos in actors.items():
        draw_actor(draw, *pos, name, actor_font)

    fill = (226, 232, 255)
    usecases = {
        "账号访问": (760, 330),
        "注册": (620, 505),
        "登录": (760, 650),
        "密码找回": (920, 505),
        "个人工作": (760, 925),
        "个人信息管理": (620, 1110),
        "加入团队": (760, 1250),
        "查看报表": (930, 1110),
        "任务协作": (1290, 925),
        "任务执行": (1135, 1110),
        "日志记录": (1290, 1250),
        "文档/代码协作": (1460, 1110),
        "团队管理": (1710, 470),
        "创建团队": (1555, 625),
        "邀请成员": (1865, 625),
        "项目管理": (1710, 880),
        "创建项目": (1555, 1040),
        "进度管控": (1865, 1040),
        "系统管理": (1710, 1270),
        "权限配置": (1455, 1400),
        "日志查看": (1710, 1400),
        "数据备份": (1965, 1400),
    }
    boxes = {name: draw_oval(draw, pos, name, fill) for name, pos in usecases.items()}

    links = {
        "访客": ["账号访问"],
        "普通用户": ["个人工作"],
        "项目成员": ["任务协作"],
        "团队管理员": ["团队管理"],
        "项目经理": ["项目管理"],
        "系统管理员": ["系统管理"],
    }
    for actor, cases in links.items():
        for case in cases:
            line_to_oval(draw, actors[actor], boxes[case])

    include_edges = [
        ("账号访问", "注册"), ("账号访问", "登录"), ("账号访问", "密码找回"),
        ("个人工作", "个人信息管理"), ("个人工作", "加入团队"), ("个人工作", "查看报表"),
        ("任务协作", "任务执行"), ("任务协作", "日志记录"), ("任务协作", "文档/代码协作"),
        ("团队管理", "创建团队"), ("团队管理", "邀请成员"),
        ("项目管理", "创建项目"), ("项目管理", "进度管控"),
        ("系统管理", "权限配置"), ("系统管理", "日志查看"), ("系统管理", "数据备份"),
    ]
    for parent, child in include_edges:
        dashed_arrow(draw, usecases[parent], nearest_point_on_oval(boxes[child], usecases[parent]), "")
    dashed_arrow(draw, usecases["登录"], nearest_point_on_oval(boxes["密码找回"], usecases["登录"]), "<<extend>>", (10, -55))
    img.save(path, quality=95)


def draw_system_diagram(path: Path):
    img = Image.new("RGB", (3200, 2100), (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title(draw, "系统用例图")
    draw.rounded_rectangle((460, 150, 2740, 1960), radius=30, fill=(255, 255, 255), outline=(14, 116, 144), width=5)
    draw.text((510, 185), "团队任务协作管理系统", fill=(21, 94, 117), font=font(34, bold=True))

    actor_font = font(27, bold=True)
    actors = {
        "访客": (210, 390),
        "普通用户": (210, 780),
        "项目成员": (210, 1260),
        "团队管理员": (2990, 500),
        "项目经理": (2990, 1030),
        "系统管理员": (2990, 1580),
    }
    for name, pos in actors.items():
        draw_actor(draw, *pos, name, actor_font)

    fill = (220, 252, 231)
    usecases = {
        "账号管理": (840, 360),
        "注册": (660, 535),
        "登录": (840, 670),
        "密码找回": (1025, 535),
        "个人工作台": (840, 910),
        "个人信息管理": (640, 1060),
        "查看任务": (840, 1210),
        "查看报表": (1045, 1060),
        "任务管理": (840, 1480),
        "新建任务": (570, 1630),
        "编辑任务": (760, 1770),
        "删除任务": (960, 1770),
        "拖拽流转": (1135, 1630),
        "添加子任务": (840, 1920),
        "日志与阻塞": (1425, 1170),
        "记录日志": (1305, 1370),
        "阻塞上报": (1550, 1370),
        "协作沟通": (1425, 1580),
        "评论": (1240, 1770),
        "@提醒": (1425, 1920),
        "上传文档": (1610, 1770),
        "代码审核": (1800, 1920),
        "团队管理": (2100, 460),
        "创建团队": (1840, 620),
        "邀请成员": (2045, 760),
        "移除成员": (2255, 760),
        "变更角色": (2460, 620),
        "项目管理": (2100, 1040),
        "创建项目": (1840, 1185),
        "配置看板": (2045, 1340),
        "分配成员": (2255, 1340),
        "进度管理": (2460, 1185),
        "阻塞处理": (2460, 1485),
        "系统管理": (2100, 1640),
        "权限管理": (1790, 1775),
        "操作日志": (1950, 1920),
        "数据备份": (2110, 1775),
        "系统参数": (2270, 1920),
        "安全策略": (2430, 1775),
    }
    boxes = {name: draw_oval(draw, pos, name, fill, width=230, height=82) for name, pos in usecases.items()}

    links = {
        "访客": ["账号管理"],
        "普通用户": ["个人工作台"],
        "项目成员": ["任务管理", "日志与阻塞", "协作沟通"],
        "团队管理员": ["团队管理"],
        "项目经理": ["项目管理", "任务管理", "日志与阻塞"],
        "系统管理员": ["系统管理"],
    }
    for actor, cases in links.items():
        for case in cases:
            line_to_oval(draw, actors[actor], boxes[case], width=2)

    include_groups = {
        "账号管理": ["注册", "登录", "密码找回"],
        "个人工作台": ["个人信息管理", "查看任务", "查看报表"],
        "任务管理": ["新建任务", "编辑任务", "删除任务", "拖拽流转", "添加子任务"],
        "日志与阻塞": ["记录日志"],
        "协作沟通": ["评论", "@提醒", "上传文档", "代码审核"],
        "团队管理": ["创建团队", "邀请成员", "移除成员", "变更角色"],
        "项目管理": ["创建项目", "配置看板", "分配成员", "进度管理"],
        "系统管理": ["权限管理", "操作日志", "数据备份", "系统参数", "安全策略"],
    }
    for parent, children in include_groups.items():
        for child in children:
            dashed_arrow(draw, usecases[parent], nearest_point_on_oval(boxes[child], usecases[parent]), "")
    dashed_arrow(draw, usecases["登录"], nearest_point_on_oval(boxes["密码找回"], usecases["登录"]), "<<extend>>", (0, -52))
    dashed_arrow(draw, usecases["阻塞上报"], nearest_point_on_oval(boxes["记录日志"], usecases["阻塞上报"]), "<<extend>>", (0, 58))
    dashed_arrow(draw, usecases["@提醒"], nearest_point_on_oval(boxes["评论"], usecases["@提醒"]), "<<extend>>", (0, 55))
    dashed_arrow(draw, usecases["阻塞处理"], nearest_point_on_oval(boxes["进度管理"], usecases["阻塞处理"]), "<<extend>>", (15, 0))

    img.save(path, quality=95)


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = paragraph._p.__class__()
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def move_paragraph_after(target, paragraph_to_move):
    target._p.addnext(paragraph_to_move._p)


def set_landscape(doc: Document):
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)


def add_image_after(doc: Document, heading_text: str, image_path: Path, caption: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == heading_text:
            caption_para = insert_paragraph_after(paragraph, caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption_para.runs:
                run.font.size = Pt(10)
                run.bold = True
            image_para = doc.add_paragraph()
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_para.add_run().add_picture(str(image_path), width=Inches(9.2))
            move_paragraph_after(caption_para, image_para)
            return
    raise ValueError(f"未找到标题：{heading_text}")


def main():
    ARTIFACTS.mkdir(parents=True, exist_ok=True)
    business = ARTIFACTS / "business-use-case.png"
    system = ARTIFACTS / "system-use-case.png"
    draw_business_diagram(business)
    draw_system_diagram(system)

    tmp = ARTIFACTS / "working.docx"
    shutil.copyfile(SRC, tmp)
    doc = Document(tmp)
    set_landscape(doc)
    add_image_after(doc, "三、业务用例图", business, "图 1 业务用例图")
    add_image_after(doc, "四、系统用例图", system, "图 2 系统用例图")
    doc.save(OUT)

    with zipfile.ZipFile(OUT) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
    print(OUT)
    print(f"media_count={len(media)}")
    print("business-use-case.png" if "business-use-case" in rels else "business image embedded via generated media name")
    print("system-use-case.png" if "system-use-case" in rels else "system image embedded via generated media name")


if __name__ == "__main__":
    main()
